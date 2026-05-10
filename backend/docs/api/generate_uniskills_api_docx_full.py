from datetime import date

from docx import Document


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Details"
    for key, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value
    doc.add_paragraph("")


def add_endpoint(doc, title, method, url, auth, request, behavior):
    doc.add_heading(title, level=3)
    rows = [
        ("Method", method),
        ("URL", url),
        ("Auth", auth),
        ("Request", request),
        ("Behavior", behavior),
    ]
    add_kv_table(doc, rows)


def main():
    doc = Document()

    doc.add_heading("UniSkills API Documentation (Full Project)", level=1)
    doc.add_paragraph(f"Generated on {date.today().isoformat()}")
    doc.add_paragraph("This document covers all HTML form endpoints in the UniSkills backend.")
    doc.add_paragraph("Base URL: /accounts/")

    doc.add_heading("Overview", level=2)
    doc.add_paragraph("The system uses Django views with HTML templates. No JSON APIs are exposed.")
    doc.add_paragraph("Roles: Student, Alumni, Admin. Staff/superuser must log in via /admin/.")

    doc.add_heading("Authentication and Registration", level=2)
    add_endpoint(
        doc,
        "Student Registration",
        "GET/POST",
        "/accounts/register/student/",
        "Public",
        "username, email, department, university_id, password1, password2",
        "Creates a student user. Email must end with @uap-bd.edu.",
    )
    add_endpoint(
        doc,
        "Alumni Registration",
        "GET/POST",
        "/accounts/register/alumni/",
        "Public",
        "username, email, graduation_year, major, current_company, password1, password2",
        "Creates an alumni user with is_alumni_verified=False.",
    )
    add_endpoint(
        doc,
        "Login",
        "GET/POST",
        "/accounts/login/",
        "Public",
        "username (or email), password",
        "Authenticates a non-admin user and redirects by role. Alumni must be verified.",
    )
    add_endpoint(
        doc,
        "Logout",
        "GET",
        "/accounts/logout/",
        "Authenticated",
        "-",
        "Logs out and redirects to landing.",
    )

    doc.add_heading("Landing and Dashboards", level=2)
    add_endpoint(
        doc,
        "Landing",
        "GET",
        "/accounts/ and /accounts/index.html",
        "Public",
        "-",
        "Landing page.",
    )
    add_endpoint(
        doc,
        "Role Redirect",
        "GET",
        "/accounts/dashboard/",
        "Authenticated",
        "-",
        "Redirects to student dashboard, alumni dashboard, or /admin/.",
    )
    add_endpoint(
        doc,
        "Student Dashboard",
        "GET",
        "/accounts/dashboard.html",
        "Authenticated",
        "-",
        "Student dashboard. Alumni are redirected.",
    )
    add_endpoint(
        doc,
        "Alumni Dashboard",
        "GET",
        "/accounts/alumni-dashboard.html",
        "Authenticated (alumni)",
        "-",
        "Alumni dashboard.",
    )

    doc.add_heading("Skill Module (Student)", level=2)
    add_endpoint(
        doc,
        "Skills Page (Browse + Create/Update)",
        "GET/POST",
        "/accounts/skills.html",
        "Authenticated",
        "GET filters: q, category, mode. POST fields: title, description, category, session_mode, available_time, fee",
        "Shows approved skill posts. Students can create/update own post (status=pending).",
    )
    add_endpoint(
        doc,
        "Skill Detail",
        "GET",
        "/accounts/skills/<post_id>/",
        "Authenticated",
        "-",
        "Shows a single approved skill post.",
    )
    add_endpoint(
        doc,
        "Delete Skill Post",
        "POST",
        "/accounts/skills/<post_id>/delete/",
        "Authenticated (owner)",
        "-",
        "Deletes the skill post.",
    )
    add_endpoint(
        doc,
        "Legacy Skill Detail",
        "GET",
        "/accounts/skill-detail.html",
        "Authenticated",
        "-",
        "Legacy route that renders the skills page.",
    )

    doc.add_heading("Skill Moderation (Admin)", level=2)
    add_endpoint(
        doc,
        "Moderation Dashboard",
        "GET",
        "/accounts/moderation.html",
        "Staff/Superuser",
        "-",
        "Shows pending skill posts.",
    )
    add_endpoint(
        doc,
        "Moderate Skill Post",
        "GET/POST",
        "/accounts/moderation/<post_id>/",
        "Staff/Superuser",
        "action=approve|reject, rejection_reason (required if reject)",
        "Approve sets status=approved. Reject sets status=rejected with reason.",
    )

    doc.add_heading("Alumni Module (UN-82/UN-83/UN-85)", level=2)
    add_endpoint(
        doc,
        "Create Alumni Post",
        "GET/POST",
        "/accounts/alumni/post/create/",
        "Verified Alumni",
        "topic, title, content, contact_link",
        "Creates AlumniPost with status=pending.",
    )
    add_endpoint(
        doc,
        "Edit Alumni Post",
        "GET/POST",
        "/accounts/alumni/post/<post_id>/edit/",
        "Author",
        "topic, title, content, contact_link",
        "Updates post and resets status=pending.",
    )
    add_endpoint(
        doc,
        "Delete Alumni Post",
        "GET/POST",
        "/accounts/alumni/post/<post_id>/delete/",
        "Author",
        "-",
        "Deletes the post after confirmation.",
    )
    add_endpoint(
        doc,
        "Browse Alumni Posts",
        "GET",
        "/accounts/alumni/posts/",
        "Public",
        "topic (optional)",
        "Shows approved alumni posts.",
    )
    add_endpoint(
        doc,
        "Alumni Directory",
        "GET",
        "/accounts/alumni/directory/",
        "Public",
        "-",
        "Lists verified alumni profiles.",
    )
    add_endpoint(
        doc,
        "Alumni Profile Detail",
        "GET",
        "/accounts/alumni/<user_id>/",
        "Public",
        "-",
        "Shows a single alumni profile and approved posts.",
    )
    add_endpoint(
        doc,
        "Alumni Moderation Dashboard",
        "GET",
        "/accounts/admin/alumni/moderation/",
        "Staff/Superuser",
        "-",
        "Shows pending alumni posts.",
    )
    add_endpoint(
        doc,
        "Approve Alumni Post",
        "GET",
        "/accounts/admin/alumni/post/<post_id>/approve/",
        "Staff/Superuser",
        "-",
        "Sets status=approved.",
    )
    add_endpoint(
        doc,
        "Reject Alumni Post",
        "GET/POST",
        "/accounts/admin/alumni/post/<post_id>/reject/",
        "Staff/Superuser",
        "reason",
        "Sets status=rejected with reason.",
    )

    doc.add_heading("Profile and Misc Pages", level=2)
    add_endpoint(
        doc,
        "Profile Update",
        "GET/POST",
        "/accounts/profile.html",
        "Authenticated",
        "username, first_name, last_name, email, department, university_id, graduation_year, major, current_company",
        "Updates profile data; student email must be @uap-bd.edu.",
    )
    add_endpoint(
        doc,
        "Rating Page",
        "GET",
        "/accounts/rating.html",
        "Authenticated",
        "-",
        "Displays rating page.",
    )
    add_endpoint(
        doc,
        "Bookings Page",
        "GET",
        "/accounts/bookings.html",
        "Authenticated",
        "-",
        "Displays bookings page.",
    )
    add_endpoint(
        doc,
        "Alumni Landing",
        "GET",
        "/accounts/alumni.html",
        "Authenticated",
        "-",
        "Displays alumni landing page.",
    )

    doc.add_heading("Admin", level=2)
    add_endpoint(
        doc,
        "Django Admin",
        "GET/POST",
        "/admin/",
        "Staff/Superuser",
        "-",
        "Admin site for managing users and posts.",
    )

    output_path = "d:/Code/Software Engineering Lab/Development UniSkills TEST/FINAL UX/backend/docs/api/UNISKILLS_API_DOCUMENTATION.docx"
    doc.save(output_path)


if __name__ == "__main__":
    main()
